import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def instructor_dashboard():
    st.title("Instructor Dashboard")

    # Initialize storage
    if "exercises" not in st.session_state:
        st.session_state["exercises"] = []
    if "translations" not in st.session_state:
        st.session_state["translations"] = []

    # --- Create New Exercise ---
    st.subheader("➕ Create New Exercise")
    new_id = len(st.session_state["exercises"]) + 1
    new_text = st.text_area("Source Text", key=f"new_text_{new_id}")
    if st.button("Add Exercise", key=f"add_ex_{new_id}"):
        if new_text.strip():
            st.session_state["exercises"].append({
                "id": new_id,
                "source_text": new_text,
                "assigned": True  # default assigned
            })
            st.success(f"Exercise {new_id} added successfully!")

    # --- Manage Existing Exercises ---
    st.subheader("📂 Existing Exercises")
    for ex in st.session_state["exercises"]:
        with st.expander(f"Exercise {ex['id']}"):
            st.write(ex["source_text"])
            ex["assigned"] = st.checkbox("Assigned", value=ex["assigned"], key=f"assigned_{ex['id']}")

    # --- Export Submissions ---
    if st.session_state["translations"]:
        st.subheader("⬇️ Download Student Submissions")
        df = pd.DataFrame(st.session_state["translations"])

        # Excel
        def to_excel(df):
            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, index=False, sheet_name="Submissions")
            return output.getvalue()

        # Word
        def to_word(df):
            output = BytesIO()
            doc = Document()
            doc.add_heading("Student Submissions", 0)
            for i, row in df.iterrows():
                doc.add_heading(f"Exercise {row['exercise_id']}", level=1)
                doc.add_paragraph(f"Source Text: {row['source_text']}")
                doc.add_paragraph(f"Translation: {row['translation']}")
                if "fluency" in row: doc.add_paragraph(f"Fluency: {row['fluency']}")
                if "accuracy" in row: doc.add_paragraph(f"Accuracy: {row['accuracy']}")
                if "time_spent_sec" in row: doc.add_paragraph(f"Time: {row['time_spent_sec']} sec")
                if "keystrokes" in row: doc.add_paragraph(f"Keystrokes: {row['keystrokes']}")
            doc.save(output)
            return output.getvalue()

        # PDF
        def to_pdf(df):
            output = BytesIO()
            doc = SimpleDocTemplate(output)
            styles = getSampleStyleSheet()
            elements = [Paragraph("Student Submissions", styles['Title']), Spacer(1,12)]
            for i, row in df.iterrows():
                elements.append(Paragraph(f"<b>Exercise {row['exercise_id']}</b>", styles['Heading2']))
                elements.append(Paragraph(f"Source Text: {row['source_text']}", styles['Normal']))
                elements.append(Paragraph(f"Translation: {row['translation']}", styles['Normal']))
                if "fluency" in row: elements.append(Paragraph(f"Fluency: {row['fluency']}", styles['Normal']))
                if "accuracy" in row: elements.append(Paragraph(f"Accuracy: {row['accuracy']}", styles['Normal']))
                if "time_spent_sec" in row: elements.append(Paragraph(f"Time Spent: {row['time_spent_sec']} sec", styles['Normal']))
                if "keystrokes" in row: elements.append(Paragraph(f"Keystrokes: {row['keystrokes']}", styles['Normal']))
                elements.append(Spacer(1,12))
            doc.build(elements)
            return output.getvalue()

        st.download_button("Download as Excel", to_excel(df), "submissions.xlsx")
        st.download_button("Download as Word", to_word(df), "submissions.docx")
        st.download_button("Download as PDF", to_pdf(df), "submissions.pdf")

# main.py  — EduApp (single file, with Localisation Lab + sticker/text/image tasks)
# Grant-oriented features:
# - Evidence-based adaptive feedback (with concrete examples)
# - Safer instructor login (env var or SHA256; explicit dev-only fallback)
# - JSON storage (no DB migration)
# - Reflection capture, per-student progress charts, class snapshot
# - Localisation Lab with configurable sticker/text/image tasks
# - Optional AI-based feedback hook for adaptive MTPE condition
# - Graceful fallbacks for optional libs; no crashes on missing deps

import os
from ai_feedback import generate_ai_feedback
import re
import json
import time
import hashlib
import random
import threading
from io import BytesIO
from pathlib import Path
from typing import List, Tuple
from difflib import SequenceMatcher, ndiff
import datetime

import streamlit as st

# Pandas is core; fail early with a clear message if missing
try:
    import pandas as pd
except Exception as e:
    st.error("This app requires pandas. Please install it with `pip install pandas`.")
    raise e

# Optional DOCX exports (graceful fallback)
try:
    from docx import Document
    from docx.shared import RGBColor
    _HAVE_DOCX = True
except Exception:
    Document = None
    RGBColor = None
    _HAVE_DOCX = False

# Optional metrics deps (graceful fallback if missing)
try:
    import sacrebleu
except Exception:
    sacrebleu = None

try:
    from bert_score import score as bertscore_score
except Exception:
    bertscore_score = None

# Optional plotting
try:
    import matplotlib.pyplot as plt  # noqa: F401
    _HAVE_MPL = True
except Exception:
    _HAVE_MPL = False

# ---------------- Proof-of-life banner (so you know this file is loaded) ----------------
THIS_FILE = os.path.abspath(__file__)
LAST_EDIT = datetime.datetime.fromtimestamp(os.path.getmtime(THIS_FILE))

# ---------------- Storage (JSON with basic locking & atomic writes) ----------------
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)

EXERCISES_FILE = DATA_DIR / "exercises.json"
SUBMISSIONS_FILE = DATA_DIR / "submissions.json"
LEADERBOARD_FILE = DATA_DIR / "leaderboard.json"

# localisation sticker/text/image tasks
LOC_STICKERS_FILE = DATA_DIR / "loc_stickers.json"
STICKER_IMG_DIR = DATA_DIR / "stickers"
STICKER_IMG_DIR.mkdir(exist_ok=True)

_lock = threading.Lock()

def load_json(file: Path):
    file = Path(file)
    if file.exists():
        with file.open("r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return {}
    return {}

def save_json(file: Path, data):
    with _lock:
        tmp = Path(str(file) + ".tmp")
        with tmp.open("w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        tmp.replace(file)

# ---------------- Auth (safer than hard-coded) ----------------
def _env(name, default=""):
    return os.getenv(name, default)

# Set one of these before running (recommended):
#   export INSTRUCTOR_PASSWORD_PLAIN='StrongPass'
#   export INSTRUCTOR_PASSWORD_SHA256='<sha256 hex of StrongPass>'
# For dev only (explicit opt-in):
#   export INSTRUCTOR_DEV_MODE='1'   # allows fallback 'admin123'
_INSTRUCTOR_PLAIN = _env("INSTRUCTOR_PASSWORD_PLAIN", "")
_INSTRUCTOR_SHA256 = _env("INSTRUCTOR_PASSWORD_SHA256", "")
_INSTRUCTOR_DEV_MODE = _env("INSTRUCTOR_DEV_MODE", "0") == "1"
_FALLBACK_PLAIN = "admin123"  # used only if DEV mode is explicitly enabled
_IS_PASSWORD_CONFIGURED = bool(_INSTRUCTOR_PLAIN or _INSTRUCTOR_SHA256 or _INSTRUCTOR_DEV_MODE)

def check_password(typed: str) -> bool:
    """
    Safer logic:
    - Prefer SHA256 env
    - Then plain env
    - Only if INSTRUCTOR_DEV_MODE=1, allow fallback 'admin123'
    - If nothing configured and no dev mode, reject all passwords
    """
    try:
        if not typed:
            return False
        if _INSTRUCTOR_SHA256:
            h = hashlib.sha256(typed.encode("utf-8")).hexdigest()
            return h == _INSTRUCTOR_SHA256
        if _INSTRUCTOR_PLAIN:
            return typed == _INSTRUCTOR_PLAIN
        if _INSTRUCTOR_DEV_MODE:
            return typed == _FALLBACK_PLAIN
        return False
    except Exception:
        return False  # never crash on login

# ---------------- Tokenization & Edit Helpers ----------------
_token_re = re.compile(r"\w+|[^\w\s]", re.UNICODE)

def _tokenize(s: str) -> List[str]:
    return _token_re.findall(s or "")

def compute_edit_details(mt_text: str, student_text: str) -> Tuple[int, int, int]:
    mt_tokens = _tokenize(mt_text)
    st_tokens = _tokenize(student_text)
    matcher = SequenceMatcher(None, mt_tokens, st_tokens)

    additions = deletions = replacements = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "insert":
            additions += (j2 - j1)
        elif tag == "delete":
            deletions += (i2 - i1)
        elif tag == "replace":
            replacements += max(i2 - i1, j2 - j1)
    total_edits = additions + deletions + replacements
    return additions, deletions, total_edits

# ---------------- Metrics ----------------
def evaluate_translation(student_text, mt_text=None, reference=None, task_type="Translate", source_text=""):
    """
    Returns a metrics dict using:
      - length_ratio (target tokens / source tokens)
      - BLEU (sacrebleu, if reference provided)
      - chrF++ (sacrebleu, if reference provided)
      - BERTScore_F1 (if available & reference provided)
      - edit counts for post-edit tasks

    All metrics gracefully fallback to None if libs or references are missing.
    """
    src_len = max(1, len(_tokenize(source_text)))
    tgt_len = len(_tokenize(student_text))
    length_ratio = round(tgt_len / src_len, 3)

    if task_type == "Post-edit MT" and mt_text:
        additions, deletions, edits = compute_edit_details(mt_text, student_text)
    else:
        additions = deletions = edits = 0

    bleu = chrf = bert_f1 = None
    if reference:
        refs = [reference]
        try:
            if sacrebleu:
                bleu = float(sacrebleu.corpus_bleu([student_text], [refs]).score)  # 0-100
                chrf = float(sacrebleu.corpus_chrf([student_text], [refs]).score)  # 0-100
        except Exception:
            bleu = None
            chrf = None
        try:
            if bertscore_score:
                P, R, F1 = bertscore_score([student_text], [reference], lang="en")
                bert_f1 = float(F1.mean().item())  # 0-1
        except Exception:
            bert_f1 = None

    return {
        "length_ratio": length_ratio,
        "BLEU": None if bleu is None else round(bleu, 2),
        "chrF++": None if chrf is None else round(chrf, 2),
        "BERTScore_F1": None if bert_f1 is None else round(bert_f1, 3),
        "additions": additions,
        "deletions": deletions,
        "edits": edits
    }

# ---------------- Track Changes (HTML + DOCX) ----------------
def _join_tokens_for_display(tokens: List[str]) -> str:
    out = " ".join(tokens)
    out = re.sub(r"\s+([.,!?;:])", r"\1", out)
    return out

def diff_text(baseline: str, student_text: str) -> str:
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    parts = []
    for w in differ:
        token = w[2:]
        if w.startswith("- "):
            parts.append(f"<span style='color:#c00;text-decoration:line-through'>{token}</span>")
        elif w.startswith("+ "):
            parts.append(f"<span style='color:#080'>{token}</span>")
        else:
            parts.append(token)
    return _join_tokens_for_display(parts)

def add_diff_to_doc(doc, baseline: str, student_text: str):
    if not _HAVE_DOCX:
        return
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    p = doc.add_paragraph()
    for w in differ:
        token = w[2:]
        if w.startswith("- "):
            run = p.add_run(token + " ")
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif w.startswith("+ "):
            run = p.add_run(token + " ")
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            p.add_run(token + " ")

# ---------------- Exports ----------------
def export_student_word(submissions, student_name):
    if not _HAVE_DOCX:
        return None
    doc = Document()
    doc.add_heading(f"Student: {student_name}", 0)
    subs = submissions.get(student_name, {})
    for ex_id, sub in subs.items():
        doc.add_heading(f"Exercise {ex_id}", level=1)
        doc.add_paragraph("Source Text:")
        doc.add_paragraph(sub.get("source_text", ""))
        if sub.get("mt_text"):
            doc.add_paragraph("MT Output:")
            doc.add_paragraph(sub.get("mt_text", ""))

        if sub.get("task_type") == "Post-edit MT":
            doc.add_paragraph("Student Submission (Track Changes):")
            base = sub.get("mt_text", "") or ""
            add_diff_to_doc(doc, base, sub.get("student_text", ""))
        else:
            doc.add_paragraph("Student Submission:")
            doc.add_paragraph(sub.get("student_text", ""))

        metrics = sub.get("metrics", {})
        doc.add_paragraph(f"Metrics: {metrics}")
        fb = sub.get("feedback")
        if fb:
            doc.add_paragraph("Feedback:")
            for line in fb:
                doc.add_paragraph(line, style="List Bullet")
        doc.add_paragraph(f"Task Type: {sub.get('task_type','')}")
        doc.add_paragraph(f"Time Spent: {sub.get('time_spent_sec', 0):.2f} sec")
        doc.add_paragraph(f"Characters (not keystrokes): {sub.get('keystrokes', 0)}")
        if sub.get("reflection"):
            doc.add_paragraph("Reflection:")
            doc.add_paragraph(sub.get("reflection"))
        if sub.get("ai_literacy_reflection"):
            doc.add_paragraph("AI Literacy Reflection:")
            doc.add_paragraph(sub.get("ai_literacy_reflection"))
        doc.add_paragraph("---")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_summary_excel(submissions):
    rows = []
    for student, subs in submissions.items():
        for ex_id, sub in subs.items():
            m = sub.get("metrics", {})
            fb = sub.get("feedback") or []
            rows.append({
                "Student": student,
                "Exercise": ex_id,
                "Task Type": sub.get("task_type", ""),
                "Condition": sub.get("condition", ""),
                "Level": sub.get("level", ""),
                "Direction": sub.get("direction", ""),
                "Length Ratio": m.get("length_ratio"),
                "BLEU": m.get("BLEU"),
                "chrF++": m.get("chrF++"),
                "BERTScore_F1": m.get("BERTScore_F1"),
                "Additions": m.get("additions"),
                "Deletions": m.get("deletions"),
                "Edits": m.get("edits"),
                "Time Spent (s)": sub.get("time_spent_sec", 0),
                "Characters Typed": sub.get("keystrokes", 0),
                "Reflection": sub.get("reflection", ""),
                "AI Literacy Reflection": sub.get("ai_literacy_reflection", ""),
                "First Feedback Item": fb[0] if fb else "",
            })
    df = pd.DataFrame(rows)
    buf = BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return buf

# Anonymised MTPE corpus export (CSV) for research
def export_anonymised_corpus(submissions):
    """
    Returns a CSV buffer with:
    - anonymised student ID (hashed)
    - exercise ID
    - task type, condition, level, direction
    - source, MT, student text
    - timing/keystrokes
    - metrics as JSON string
    - counts of AI interactions
    """
    rows = []
    for student, subs in submissions.items():
        anon_id = hashlib.sha256(student.encode("utf-8")).hexdigest()[:10]
        for ex_id, sub in subs.items():
            metrics = sub.get("metrics", {})
            ai_interactions = sub.get("ai_interactions", []) or []
            rows.append({
                "AnonID": anon_id,
                "Exercise": ex_id,
                "TaskType": sub.get("task_type", ""),
                "Condition": sub.get("condition", ""),
                "Level": sub.get("level", ""),
                "Direction": sub.get("direction", ""),
                "SourceText": sub.get("source_text", ""),
                "MTText": sub.get("mt_text", ""),
                "StudentText": sub.get("student_text", ""),
                "TimeSpent_sec": sub.get("time_spent_sec", 0),
                "CharactersTyped": sub.get("keystrokes", 0),
                "MetricsJSON": json.dumps(metrics, ensure_ascii=False),
                "AIInteractionCount": len(ai_interactions),
            })
    df = pd.DataFrame(rows)
    buf = BytesIO()
    df.to_csv(buf, index=False)
    buf.seek(0)
    return buf

# ---------------- Gamification ----------------
def load_leaderboard():
    return load_json(LEADERBOARD_FILE)

def update_leaderboard(student_name, points):
    leaderboard = load_leaderboard()
    leaderboard[student_name] = leaderboard.get(student_name, 0) + points
    save_json(LEADERBOARD_FILE, leaderboard)

def show_leaderboard():
    leaderboard = load_leaderboard()
    st.subheader("Leaderboard")
    if leaderboard:
        items = sorted(leaderboard.items(), key=lambda x: x[1], reverse=True)
        df = pd.DataFrame(items, columns=["Student", "Points"])
        st.dataframe(df, use_container_width=True)
    else:
        st.info("No leaderboard data yet.")

# ---------------- Instructor: manage sticker/text/image localisation tasks ----------------
def localisation_sticker_manager():
    st.subheader("Instructor – Sticker / text / image localisation tasks")

    if not _IS_PASSWORD_CONFIGURED:
        st.info(
            "Instructor password is not configured. "
            "Set INSTRUCTOR_PASSWORD_PLAIN or INSTRUCTOR_PASSWORD_SHA256 "
            "(optionally INSTRUCTOR_DEV_MODE=1 for a dev fallback)."
        )

    pwd = st.text_input("Instructor password", type="password", key="loc_sticker_pwd")
    if not check_password(pwd):
        st.info("Enter a valid instructor password to manage these tasks.")
        return

    loc_stickers = load_json(LOC_STICKERS_FILE)
    sticker_ids = ["New task"] + sorted(loc_stickers.keys())
    selection = st.selectbox("Choose task", sticker_ids, key="loc_sticker_select")

    if selection != "New task" and selection in loc_stickers:
        current = loc_stickers[selection]
        default_title = current.get("title", "")
        default_instr = current.get("instructions", "")
        default_text = current.get("content_text", "")
        default_url = current.get("image_url", "") if current.get("image_type") == "url" else ""

        st.markdown("**Current preview for students:**")
        if default_text:
            st.markdown("**Text to localise:**")
            st.write(default_text)
        if current.get("image_type") == "uploaded":
            img_path = current.get("image_path", "")
            if img_path and Path(img_path).exists():
                st.image(str(img_path))
            else:
                st.warning("Image file not found on server.")
        elif current.get("image_type") == "url":
            st.image(current.get("image_url", ""))
    else:
        default_title = ""
        default_instr = ""
        default_text = ""
        default_url = ""

    with st.form("loc_sticker_form"):
        title = st.text_input("Task title", value=default_title)
        content_text = st.text_area(
            "Text to be localised (optional)",
            value=default_text,
            height=120
        )
        instructions = st.text_area(
            "Instructions for students (what to do with this text / image)",
            value=default_instr,
            height=120
        )

        st.write("Sticker / image (choose either URL or upload, both optional):")
        image_url = st.text_input("Image URL (optional)", value=default_url)
        uploaded = st.file_uploader(
            "Or upload an image file",
            type=["png", "jpg", "jpeg", "webp"]
        )

        col1, col2 = st.columns(2)
        with col1:
            save_btn = st.form_submit_button("Save / Update task")
        with col2:
            delete_btn = st.form_submit_button("Delete this task")

    if save_btn:
        if not title.strip() and not content_text.strip() and not image_url and not uploaded:
            st.warning("Please provide at least a title, some text, or an image before saving.")
            return

        # choose ID
        if selection != "New task" and selection in loc_stickers:
            task_id = selection
        else:
            existing_nums = []
            for sid in loc_stickers.keys():
                m = re.match(r"STK_(\d+)$", sid)
                if m:
                    existing_nums.append(int(m.group(1)))
            next_num = max(existing_nums + [0]) + 1
            task_id = f"STK_{next_num:03d}"

        image_type = None
        image_path = ""
        image_url_final = ""

        if uploaded is not None:
            safe_name = re.sub(r"[^\w\.\-]", "_", uploaded.name)
            file_path = STICKER_IMG_DIR / f"{task_id}_{safe_name}"
            with open(file_path, "wb") as f:
                f.write(uploaded.getbuffer())
            image_type = "uploaded"
            image_path = str(file_path)
        elif image_url:
            image_type = "url"
            image_url_final = image_url

        # if updating and no new image chosen, keep old image info
        if task_id in loc_stickers and image_type is None:
            image_type = loc_stickers[task_id].get("image_type")
            image_path = loc_stickers[task_id].get("image_path", "")
            image_url_final = loc_stickers[task_id].get("image_url", "")

        loc_stickers[task_id] = {
            "title": title.strip(),
            "instructions": instructions.strip(),
            "content_text": content_text.strip(),
            "image_type": image_type,
            "image_path": image_path,
            "image_url": image_url_final,
            "created_at": datetime.datetime.now().isoformat()
        }
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {task_id} saved.")

    if delete_btn and selection != "New task" and selection in loc_stickers:
        loc_stickers.pop(selection, None)
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {selection} deleted.")

# ---------------- Optional AI generator & feedback helper ----------------
def ai_generate_text(prompt: str):
    """
    Generic text generation hook.
    Currently uses HuggingFace Inference API if HF_API_TOKEN is set.
    If not configured or failing, returns None (graceful fallback).
    """
    HF_TOKEN = os.getenv("HF_API_TOKEN", "") or st.secrets.get("HF_API_TOKEN", "")
    if not HF_TOKEN:
        return None
    try:
        import requests
        headers = {"Authorization": f"Bearer {HF_TOKEN}"}
        payload = {"inputs": prompt}
        response = requests.post(
            "https://api-inference.huggingface.co/models/gpt2",
            headers=headers,
            json=payload,
            timeout=15
        )
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list) and data and "generated_text" in data[0]:
                return data[0]["generated_text"]
    except Exception:
        pass
    return None

def get_ai_feedback_on_text(source_text: str, mt_text: str, student_text: str, task_type: str):
    """
    Wraps ai_generate_text with a prompt geared towards:
    - accuracy
    - register
    - idiomatic fidelity
    - stylistic appropriateness

    Returns a string (feedback/suggestions) or None if no backend configured.
    """
    base_prompt = """
You are an expert English–Arabic translation trainer focusing on Machine Translation Post-Editing (MTPE).
Analyse the student's current version in relation to the source (and MT output if given).
Comment briefly (bullet points) on:

1. Accuracy (meaning preservation)
2. Register (formality, appropriateness for context)
3. Idiomatic expressions and culturally bound elements
4. Overall coherence and style

Then, if helpful, suggest ONE revised version of 1–2 sentences that illustrates better choices,
but keep your answer under 250 words.

Use Arabic where appropriate in examples.
"""
    parts = [
        base_prompt.strip(),
        "\n\n=== SOURCE TEXT ===\n",
        source_text or "(none)",
        "\n\n=== MT OUTPUT (if any) ===\n",
        mt_text or "(none)",
        "\n\n=== STUDENT VERSION ===\n",
        student_text or "(empty)",
        f"\n\nTask type: {task_type}\n"
    ]
    prompt = "".join(parts)
    return ai_generate_text(prompt)

# ---------------- Evidence-based Linguistic Hints ----------------
_AR_LETTERS = r"\u0600-\u06FF"  # Arabic Unicode block

def _tokenize_words(text: str):
    # words incl. hyphen/apostrophes; keep numbers as tokens
    return re.findall(r"[A-Za-z" + _AR_LETTERS + r"]+[’'\-]?[A-Za-z" + _AR_LETTERS + r"]+|\d+(?:[.,]\d+)?", text)

def _likely_terms(source_text: str):
    """
    Heuristics for 'terms/proper names':
    - Titlecase/ALLCAPS (Latin)
    - Contains hyphen or digits
    - Quoted spans
    - Arabic words length>=4
    """
    terms = set()
    # quoted chunks
    for q in re.findall(r"[\"“”‘’'`«»](.+?)[\"“”‘’'`«»]", source_text):
        for w in _tokenize_words(q):
            if len(w) >= 3:
                terms.add(w)

    for w in _tokenize_words(source_text):
        if re.match(r"[A-Z][A-Za-z\-]+$", w):          # Titlecase
            terms.add(w)
        elif re.match(r"[A-Z0-9\-]{3,}$", w):          # ALLCAPS or alnum with -
            terms.add(w)
        elif "-" in w or re.search(r"\d", w):          # hyphenated or digits
            terms.add(w)
        elif re.match(r"[" + _AR_LETTERS + r"]{4,}$", w):  # Arabic word len>=4
            terms.add(w)
    return terms

def _short_list(items, n=4):
    items = list(items)
    if not items:
        return ""
    if len(items) <= n:
        return " | ".join(items)
    return " | ".join(items[:n]) + f" … (+{len(items)-n} more)"

def quick_linguistic_hints(source_text: str, student_text: str):
    hints = []
    try:
        # Numbers: exact evidence
        src_nums = set(re.findall(r"\d+(?:[.,]\d+)?", source_text))
        tgt_nums = set(re.findall(r"\d+(?:[.,]\d+)?", student_text))
        missing_nums = sorted(src_nums - tgt_nums, key=lambda x: (len(x), x))
        if missing_nums:
            hints.append({
                "rule": "numbers_missing",
                "message": "Some figures from the source didn’t appear in your text.",
                "evidence": f"Missing: {_short_list(missing_nums)}"
            })

        # Brackets & quotes balance
        for sym_open, sym_close, label in [("(", ")", "parentheses"), ("[", "]", "brackets"), ("{", "}", "braces")]:
            if source_text.count(sym_open) != student_text.count(sym_close):
                hints.append({
                    "rule": f"{label}_unbalanced",
                    "message": f"{label.capitalize()} look unbalanced.",
                    "evidence": (f"Source {sym_open}/{sym_close}: {source_text.count(sym_open)}/{source_text.count(sym_close)}; "
                                 f"Your text: {student_text.count(sym_open)}/{student_text.count(sym_close)}")
                })
        if source_text.count('"') != student_text.count('"'):
            hints.append({
                "rule": "quotes_unbalanced",
                "message": "Quotation marks may be unbalanced.",
                "evidence": f'Source quotes: {source_text.count(chr(34))}; Yours: {student_text.count(chr(34))}'
            })

        # Terms/proper names: concrete examples
        src_terms = _likely_terms(source_text)
        tgt_tokens = set(_tokenize_words(student_text))
        missing_terms = sorted([t for t in src_terms if t not in tgt_tokens], key=lambda x: (-len(x), x))
        if missing_terms:
            hints.append({
                "rule": "terms_missing",
                "message": "Some key terms/names from the source weren’t reflected.",
                "evidence": f"Examples: {_short_list(missing_terms)}"
            })
    except Exception:
        pass
    return hints

# ---------------- Adaptive Feedback ----------------
def generate_feedback(metrics: dict, task_type: str, source_text: str, student_text: str, extra_hints=None):
    """
    Return a list of bullet strings that explicitly mention the metric thresholds/evidence.
    """
    msgs = []
    lr = metrics.get("length_ratio")
    edits = int(metrics.get("edits", 0) or 0)
    adds = int(metrics.get("additions", 0) or 0)
    dels = int(metrics.get("deletions", 0) or 0)
    bleu = metrics.get("BLEU")
    chrf = metrics.get("chrF++")

    # 1) Edit profile (Post-edit MT)
    if task_type == "Post-edit MT":
        if edits == 0:
            msgs.append(("edits_none",
                         "No edits were applied to the MT output.",
                         "Edits = 0; review the MT carefully—critical errors may remain."))
        elif edits > 20:
            msgs.append(("edits_many",
                         f"High edit volume detected: {edits} edits (additions {adds}, deletions {dels}).",
                         "Focus on adequacy first; too many surface edits can hide missed meaning issues."))

    # 2) Length ratio diagnostics
    if lr is not None:
        if lr < 0.80:
            msgs.append(("len_low",
                         f"Your length ratio is {lr:.2f} (target ≈ 0.90–1.20).",
                         "This suggests compression/omissions; check if any propositions or modifiers were dropped."))
        elif lr > 1.30:
            msgs.append(("len_high",
                         f"Your length ratio is {lr:.2f} (target ≈ 0.90–1.20).",
                         "Text may be over-expanded; look for repetition or overly literal padding."))

    # 3) Metric interplay (accuracy vs fluency)
    if bleu is not None and chrf is not None:
        if bleu < 30 <= chrf:
            msgs.append(("acc_low_flu_ok",
                         f"chrF++ is {chrf:.1f} (character-level similarity ok) but BLEU is {bleu:.1f} (segment overlap low).",
                         "Terminology/lexical choices may diverge from a typical solution—revisit key terms and content words."))
        elif bleu >= 30 and chrf < 50:
            msgs.append(("flu_low_acc_ok",
                         f"BLEU is {bleu:.1f} (content roughly aligned) but chrF++ is {chrf:.1f} (lower fluency).",
                         "Try smoothing longer clauses, connectors, and sentence rhythm to improve readability."))
        elif bleu is not None and bleu < 20:
            msgs.append(("both_low",
                         f"BLEU is {bleu:.1f}, which is in a low range.",
                         "Start by checking whether every idea in the source is present before polishing style."))

    # 4) Integrate extra hints (numbers/terms/quotes) with evidence
    if extra_hints:
        for h in extra_hints:
            rule = h.get("rule", "hint")
            msg = h.get("message", "")
            evd = h.get("evidence", "")
            if evd:
                msgs.append((rule, msg, evd))
            else:
                msgs.append((rule, msg, ""))

    # De-duplicate by rule key, keep order, cap to top 4
    seen = set()
    final = []
    for key, text, detail in msgs:
        if key in seen:
            continue
        seen.add(key)
        if detail:
            final.append(f"• {text} — *{detail}*")
        else:
            final.append(f"• {text}")
        if len(final) >= 4:
            break
    return final

# ---------------- Instructor Dashboard (Core Translation Lab) ----------------
def instructor_dashboard():
    st.title("Instructor Dashboard")

    if not _IS_PASSWORD_CONFIGURED:
        st.info(
            "Instructor password is not configured. "
            "Set INSTRUCTOR_PASSWORD_PLAIN or INSTRUCTOR_PASSWORD_SHA256 for production use. "
            "Optionally, use INSTRUCTOR_DEV_MODE=1 with fallback 'admin123' in a dev environment only."
        )

    password = st.text_input("Enter instructor password", type="password")
    if not check_password(password):
        st.warning("Incorrect password or password not configured. Access denied.")
        return

    exercises = load_json(EXERCISES_FILE)
    submissions = load_json(SUBMISSIONS_FILE)

    st.subheader("Create / Edit / Delete Exercise")
    ex_ids = ["New"] + list(exercises.keys())
    selected_ex = st.selectbox("Select Exercise", ex_ids)

    # Prefill if editing
    if selected_ex != "New" and selected_ex in exercises:
        default_source = exercises[selected_ex].get("source_text", "")
        default_mt = exercises[selected_ex].get("mt_text", "") or ""
    else:
        default_source = ""
        default_mt = ""

    with st.form("exercise_form"):
        st_text = st.text_area("Source Text", value=default_source, height=150)
        mt_text = st.text_area("MT Output (optional)", value=default_mt, height=150)
        col1, col2, col3 = st.columns(3)
        with col1:
            save_btn = st.form_submit_button("Save Exercise")
        with col2:
            delete_btn = st.form_submit_button("Delete Exercise")
        with col3:
            gen_btn = st.form_submit_button("Generate AI Exercise")

    if save_btn:
        try:
            next_id = (
                str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
                if selected_ex == "New" else selected_ex
            )
        except Exception:
            next_id = "001" if selected_ex == "New" else selected_ex

        exercises[next_id] = {
            "source_text": st_text,
            "mt_text": (mt_text.strip() or None)
        }
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved! ID: {next_id}")

    if delete_btn and selected_ex != "New":
        exercises.pop(selected_ex, None)
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise {selected_ex} deleted!")

    if gen_btn:
        prompt = "Write a short culturally rich text for translation students."
        ai_text = ai_generate_text(prompt)
        new_text = ai_text if ai_text else f"This is AI generated exercise {random.randint(1,1000)}."
        new_mt = f"MT output for exercise {random.randint(1,1000)}."
        try:
            next_id = str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
        except Exception:
            next_id = "001"
        exercises[next_id] = {"source_text": new_text, "mt_text": new_mt}
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved as ID {next_id}")

    st.subheader("Download Exercises")
    if exercises and _HAVE_DOCX:
        for ex_id, ex in exercises.items():
            try:
                buf = BytesIO()
                doc = Document()
                doc.add_heading(f"Exercise {ex_id}", 0)
                doc.add_paragraph("Source Text:")
                doc.add_paragraph(ex.get("source_text", ""))
                if ex.get("mt_text"):
                    doc.add_paragraph("MT Output:")
                    doc.add_paragraph(ex.get("mt_text", ""))
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    f"Exercise {ex_id} (Word)",
                    buf,
                    file_name=f"Exercise_{ex_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception:
                st.info(f"Exercise {ex_id}: export not available (DOCX error).")
    elif exercises and not _HAVE_DOCX:
        st.info("Word export is disabled because python-docx is not installed.")
    else:
        st.info("No exercises yet.")

    st.subheader("Student Submissions & Exports")
    if submissions:
        student_choice = st.selectbox("Choose student", ["All"] + list(submissions.keys()))
        if student_choice != "All":
            buf = export_student_word(submissions, student_choice)
            if buf is not None:
                safe_name = re.sub(r"[^\w\-]+", "_", student_choice)
                st.download_button(
                    f"Download {student_choice}'s Submissions (Word)",
                    buf,
                    file_name=f"{safe_name}_submissions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.info("Word export unavailable (python-docx not installed).")

        st.subheader("Download Metrics Summary (Excel)")
        excel_buf = export_summary_excel(submissions)
        st.download_button(
            "Download Excel Summary",
            excel_buf,
            file_name="metrics_summary.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.subheader("Download anonymised MTPE corpus (CSV)")
        anon_buf = export_anonymised_corpus(submissions)
        st.download_button(
            "Download anonymised corpus",
            anon_buf,
            file_name="mtpe_corpus_anon.csv",
            mime="text/csv"
        )

        # Class Snapshot (mean chrF++ per exercise)
        try:
            st.subheader("Class Snapshot (chrF++)")
            rows = []
            for ex_id2, ex in exercises.items():
                vals = []
                for student, subs in submissions.items():
                    sub = subs.get(ex_id2)
                    if sub:
                        m = sub.get("metrics", {})
                        if m.get("chrF++") is not None:
                            vals.append(m["chrF++"])
                if vals:
                    mean_val = round(sum(vals) / max(1, len(vals)), 2)
                    rows.append({"Exercise": ex_id2, "chrF++ mean": mean_val, "n": len(vals)})
            if rows:
                df_snapshot = pd.DataFrame(rows)
                st.dataframe(df_snapshot, use_container_width=True)
            else:
                st.info("No metrics yet to summarize.")
        except Exception:
            st.info("Snapshot unavailable (aggregation error).")

        # Reflection snapshot
        try:
            st.subheader("Reflection Snapshot (by exercise)")
            refl_rows = []
            for student, subs in submissions.items():
                for ex_id2, sub in subs.items():
                    refl = (sub.get("reflection") or "").strip()
                    if refl:
                        refl_rows.append({
                            "Exercise": ex_id2,
                            "Student": student,
                            "Reflection (snippet)": (refl[:120] + "…") if len(refl) > 120 else refl
                        })
            if refl_rows:
                st.dataframe(pd.DataFrame(refl_rows), use_container_width=True)
            else:
                st.info("No reflections recorded yet.")
        except Exception:
            st.info("Reflection snapshot unavailable.")

        show_leaderboard()
    else:
        st.info("No submissions yet.")

# ---------------- Student Dashboard (Core Translation Lab) ----------------
def student_dashboard():
    st.title("Student Dashboard")

    exercises = load_json(EXERCISES_FILE)
    if not exercises:
        st.info("No exercises available yet. Please check back later.")
        return

    submissions = load_json(SUBMISSIONS_FILE)

    student_name = st.text_input("Enter your name or student ID")
    if not student_name:
        return

    # Course level + direction (for research covariates)
    col_lvl, col_dir = st.columns(2)
    with col_lvl:
        level = st.selectbox(
            "Course level (for research use)",
            ["Not specified", "Beginner", "Intermediate", "Advanced"],
            index=0
        )
    with col_dir:
        direction = st.selectbox(
            "Translation direction",
            ["Not specified", "EN → AR", "AR → EN"],
            index=0
        )

    # Study condition (adaptive vs traditional)
    condition = st.radio(
        "Study condition (set by your instructor)",
        ["Traditional MTPE", "Adaptive MTPE (AI-assisted)"],
        index=0,
        help=(
            "Traditional MTPE = no AI suggestions are used for this task.\n"
            "Adaptive MTPE = you may request AI feedback/suggestions during post-editing."
        )
    )

    # Ensure per-student container in JSON
    if student_name not in submissions:
        submissions[student_name] = {}

    ex_id = st.selectbox("Choose Exercise", list(exercises.keys()))
    if not ex_id:
        return

    ex = exercises[ex_id]
    st.subheader("Source Text")
    st.markdown(
        f"<div style='font-family:Times New Roman;font-size:12pt;'>{ex.get('source_text','')}</div>",
        unsafe_allow_html=True
    )

    task_options = ["Translate"] if not ex.get("mt_text") else ["Translate", "Post-edit MT"]
    task_type = st.radio("Task Type", task_options, horizontal=True)

    initial_text = "" if task_type == "Translate" else (ex.get("mt_text", "") or "")

    # Session keys unique per student/exercise
    start_key = f"start_time_{student_name}_{ex_id}"
    keys_key = f"chars_{student_name}_{ex_id}"
    ai_log_key = f"ai_log_{student_name}_{ex_id}"

    if start_key not in st.session_state:
        st.session_state[start_key] = time.time()
    if keys_key not in st.session_state:
        st.session_state[keys_key] = 0
    if ai_log_key not in st.session_state:
        st.session_state[ai_log_key] = []

    with st.form(key=f"exercise_form_{student_name}_{ex_id}"):
        student_text = st.text_area(
            "Type your translation / post-edit here",
            initial_text,
            height=300
        )
        reflection = st.text_area(
            "Brief reflection (what changed / why?)",
            "",
            height=80
        )
        ai_lit_reflection = st.text_area(
            "AI literacy reflection (optional):\n"
            "- If you used MT/AI, did you notice hallucinations, bias, register problems, or copyright issues?\n"
            "- How did you decide when to trust or override the machine output?",
            "",
            height=120
        )

        col_submit, col_ai = st.columns(2)
        with col_submit:
            submitted = st.form_submit_button("Submit")
        with col_ai:
            ask_ai = st.form_submit_button("Get AI feedback / suggestion")

    # Handle AI feedback request (only in Adaptive condition)
  if ask_ai:
       if condition.startswith("Adaptive"):

        prompt = f"""
You are an expert English–Arabic translation trainer.

SOURCE:
{ex.get("source_text", "")}

MT OUTPUT:
{ex.get("mt_text", "")}

STUDENT VERSION:
{student_text}

Give concise bullet-point feedback on:
1. Accuracy
2. Register
3. Idiomatic use
4. Style
"""

        ai_feedback = generate_ai_feedback(prompt)

        if ai_feedback:
            st.markdown("### AI feedback / suggestion")
            st.write(ai_feedback)

            st.session_state[ai_log_key].append({
                "timestamp": datetime.datetime.now().isoformat(),
                "condition": condition,
                "task_type": task_type,
                "student_text": student_text,
                "ai_feedback": ai_feedback,
            })

        else:
            st.warning(
                "The AI feedback service could not be reached. "
                "Check HF_API_TOKEN or your internet connection."
            )

    else:
        st.info(
            "You are in the 'Traditional MTPE' condition. "
            "AI suggestions are disabled for this task by design."
        )
    if submitted:
        time_spent = time.time() - st.session_state[start_key]
        st.session_state[keys_key] = len(student_text)  # characters typed proxy

        metrics = evaluate_translation(
            student_text,
            mt_text=ex.get("mt_text"),
            reference=None,
            task_type=task_type,
            source_text=ex.get("source_text", "")
        )

        extra = quick_linguistic_hints(ex.get("source_text",""), student_text)
        feedback_msgs = generate_feedback(
            metrics,
            task_type,
            ex.get("source_text",""),
            student_text,
            extra
        )

        # Persist submission (including research fields and AI interactions)
        submissions[student_name][ex_id] = {
            "source_text": ex.get("source_text", ""),
            "mt_text": ex.get("mt_text"),
            "student_text": student_text,
            "task_type": task_type,
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": st.session_state[keys_key],
            "metrics": metrics,
            "reflection": reflection,
            "ai_literacy_reflection": ai_lit_reflection,
            "feedback": feedback_msgs,
            "condition": condition,
            "level": level,
            "direction": direction,
            "ai_interactions": st.session_state[ai_log_key],
        }
        save_json(SUBMISSIONS_FILE, submissions)

        # Gamification points (BLEU/chrF++ might be None if no reference)
        points = 0
        try:
            if metrics.get("BLEU") is not None:
                points += int(metrics["BLEU"])
            if metrics.get("chrF++") is not None:
                points += int(metrics["chrF++"] / 2)
            if task_type == "Post-edit MT":
                points += max(0, 10 - int(metrics["edits"]))
        except Exception:
            pass
        update_leaderboard(student_name, points)

        st.success("Submission saved!")

        # Show metrics neatly
        def _fmt(v):
            return "—" if v is None else v
        st.subheader("Your Metrics")
        st.markdown(f"""
- **Length Ratio** (target/src): {_fmt(metrics['length_ratio'])}
- **BLEU**: {_fmt(metrics['BLEU'])}
- **chrF++**: {_fmt(metrics['chrF++'])}
- **BERTScore F1**: {_fmt(metrics['BERTScore_F1'])}
- **Additions**: {_fmt(metrics['additions'])}
- **Deletions**: {_fmt(metrics['deletions'])}
- **Edits**: {_fmt(metrics['edits'])}
- **Time Spent**: {round(time_spent, 2)} sec
- **Characters Typed**: {st.session_state[keys_key]}
- **Condition**: {condition}
""")

        # Adaptive feedback
        st.subheader("Adaptive Feedback")
        if feedback_msgs:
            for m in feedback_msgs:
                st.markdown(m)
        else:
            st.info("No specific issues triggered. Focus on cohesion, clarity, and consistent terminology.")

        if task_type == "Post-edit MT":
            st.subheader("Track Changes")
            st.caption("Track changes: green = additions, red strike = deletions.")
            base = ex.get("mt_text", "") or ""
            st.markdown(diff_text(base, student_text), unsafe_allow_html=True)

        # Progress mini-dashboard (JSON-based)
        try:
            history = []
            for ex_id2, sub2 in submissions.get(student_name, {}).items():
                m2 = sub2.get("metrics", {})
                history.append({
                    "ex": ex_id2,
                    "BLEU": m2.get("BLEU"),
                    "chrF++": m2.get("chrF++"),
                    "Edits": m2.get("edits", 0)
                })
            if history:
                st.subheader("Progress Overview")
                df_hist = pd.DataFrame(history)
                try:
                    if not df_hist.empty:
                        df_trend = df_hist.set_index("ex")[["BLEU","chrF++"]]
                        st.line_chart(df_trend)
                except Exception:
                    pass
                try:
                    df_edits = df_hist.set_index("ex")[["Edits"]]
                    st.bar_chart(df_edits)
                except Exception:
                    pass
        except Exception:
            st.info("Progress charts unavailable.")

        show_leaderboard()

# ---------------- Localisation Lab (with JSON + leaderboard + stickers/text/images) ----------------
def localisation_lab():
    st.title("🌍 Localisation Lab")
    st.write(
        "Interactive exercises on localisation (English ↔ Arabic). "
        "Work here is saved to the same JSON/leaderboard as the core lab."
    )

    mode = st.sidebar.radio(
        "Localisation mode",
        ["Student view", "Instructor (manage sticker/text/image tasks)"],
        index=0,
        key="loc_mode"
    )

    if mode == "Instructor (manage sticker/text/image tasks)":
        localisation_sticker_manager()
        return

    # --- student view from here on ---

    # Identify student so we can save work
    student_name = st.text_input("Enter your name (for saving localisation work)")
    if not student_name:
        st.info("Please enter your name to start.")
        return

    submissions = load_json(SUBMISSIONS_FILE)
    if student_name not in submissions:
        submissions[student_name] = {}

    exercise = st.sidebar.selectbox(
        "Choose a localisation exercise",
        [
            "1️⃣ Translation vs Localisation",
            "2️⃣ Cultural Adaptation in Advertising",
            "3️⃣ Conventions: Dates, Units, Currency",
            "4️⃣ Tone & Website/App UX",
            "5️⃣ Post-editing: Error Detection",
            "6️⃣ App Store Description",
            "7️⃣ Strategy & Theory Reflection",
            "🎨 Sticker / text / image task (from instructor)",
        ],
        key="loc_ex_select",
    )

    # Helper to save a localisation submission and show feedback/metrics
    def save_loc_submission(
        ex_id: str,
        source_text: str,
        main_text: str,
        reflection_text: str,
    ):
        if not main_text.strip():
            st.warning("Nothing to save yet — please write your main answer first.")
            return

        # Time + 'keystrokes'
        start_key = f"loc_start_{student_name}_{ex_id}"
        if start_key not in st.session_state:
            st.session_state[start_key] = time.time()
        time_spent = time.time() - st.session_state[start_key]

        keystrokes = len(main_text)

        metrics = evaluate_translation(
            main_text,
            mt_text=None,
            reference=None,
            task_type="Localisation",
            source_text=source_text,
        )

        extra = quick_linguistic_hints(source_text, main_text)
        feedback_msgs = generate_feedback(
            metrics,
            "Localisation",
            source_text,
            main_text,
            extra_hints=extra,
        )

        # Build submission record
        submissions[student_name][ex_id] = {
            "source_text": source_text,
            "mt_text": None,
            "student_text": main_text,
            "task_type": "Localisation",
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": keystrokes,
            "metrics": metrics,
            "reflection": reflection_text,
            "feedback": feedback_msgs,
        }
        save_json(SUBMISSIONS_FILE, submissions)

        # Simple participation-based points + small bonus for reasonable length ratio
        points = 15
        lr = metrics.get("length_ratio")
        try:
            if lr is not None and 0.8 <= lr <= 1.3:
                points += 5
        except Exception:
            pass
        update_leaderboard(student_name, points)

        st.success("Localisation submission saved and leaderboard updated!")

        # Show metrics (no tricky multi-line f-string)
        def _fmt(v):
            return "—" if v is None else v

        st.subheader("Your Metrics (Localisation)")
        st.write(f"• Length Ratio (target/src): {_fmt(metrics['length_ratio'])}")
        st.write(f"• BLEU: {_fmt(metrics['BLEU'])}")
        st.write(f"• chrF++: {_fmt(metrics['chrF++'])}")
        st.write(f"• BERTScore F1: {_fmt(metrics['BERTScore_F1'])}")
        st.write(f"• Time Spent: {round(time_spent, 2)} sec")
        st.write(f"• Characters Typed: {keystrokes}")

        st.subheader("Adaptive Feedback")
        if feedback_msgs:
            for m in feedback_msgs:
                st.markdown(m)
        else:
            st.info("No specific issues triggered. Focus on cohesion, clarity, and consistent localisation choices.")

        st.subheader("Leaderboard (including localisation tasks)")
        show_leaderboard()

    # ---- Text-based exercises (1–7) ----

    def exercise_1():
        ex_id = "LOC_1"
        source_text = (
            "Download our app today and enjoy free shipping on all orders over $50. "
            "Offer valid through July 4. Call 1-800-555-0199 for assistance."
        )

        st.header("1️⃣ Translation vs Localisation")

        with st.form(f"loc_form_{ex_id}"):
            st.subheader("Source Text (English → Arabic)")
            st.markdown(f"> {source_text}")

            st.markdown("### Step 1 – Translation")
            literal = st.text_area(
                "Write your **initial translation into Arabic** (before localisation):",
                key="loc_ex1_translation",
                height=140,
            )

            st.markdown("### Step 2 – Identify Localisation Elements")
            elements = st.text_area(
                "List at least **5 elements** that require localisation "
                "(e.g., currency, dates, phone formats, cultural references):",
                key="loc_ex1_elements",
                height=120,
            )

            st.markdown("### Step 3 – Market-specific Localisation")
            col1, col2 = st.columns(2)
            with col1:
                uae = st.text_area(
                    "Write a **localised version for the UAE**:",
                    key="loc_ex1_uae",
                    height=160,
                )
            with col2:
                ksa = st.text_area(
                    "Write a **localised version for Saudi Arabia**:",
                    key="loc_ex1_ksa",
                    height=160,
                )

            st.markdown("### Step 4 – Reflection")
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                ling = st.text_area("**Linguistic changes**", key="loc_ex1_ling", height=100)
            with col_b:
                cult = st.text_area("**Cultural changes**", key="loc_ex1_cult", height=100)
            with col_c:
                func = st.text_area("**Functional/technical changes**", key="loc_ex1_func", height=100)

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / discussion points"):
            st.markdown(
                "- Currency (USD → AED/SAR) and thresholds\n"
                "- Date format and local holidays vs July 4\n"
                "- Phone number format and local support channels\n"
                "- Register and marketing tone in Arabic\n"
                "- App store and e-commerce conventions in UAE/KSA"
            )

        if submitted:
            reflection = (
                "Elements needing localisation:\n" + elements.strip() + "\n\n"
                "Linguistic changes:\n" + ling.strip() + "\n\n"
                "Cultural changes:\n" + cult.strip() + "\n\n"
                "Functional/technical changes:\n" + func.strip()
            )
            # Use UAE version as the main 'evaluated' text
            save_loc_submission(ex_id, source_text, uae, reflection)

    def exercise_2():
        ex_id = "LOC_2"
        source_text = (
            "Celebrate Black Friday with unbelievable deals! "
            "Grab your favorite winter outfits before the snow hits!"
        )

        st.header("2️⃣ Cultural Adaptation in Advertising")

        with st.form(f"loc_form_{ex_id}"):
            st.subheader("Source Text (English → Arabic)")
            st.markdown(f"> {source_text}")

            st.markdown("### Step 1 – Literal Translation")
            literal = st.text_area(
                "Write a **literal translation into Arabic**:",
                key="loc_ex2_literal",
                height=140,
            )

            st.markdown("### Step 2 – Localised Version for a Gulf Audience")
            gulf = st.text_area(
                "Now write a **localised version for a Gulf audience**:",
                key="loc_ex2_localised",
                height=160,
            )

            st.markdown("### Step 3 – Strategic Choices")
            bf_choice = st.radio(
                "What do you do with **“Black Friday”**?",
                [
                    "Keep as Black Friday (in English or transliterated)",
                    "Use an existing local term (e.g. White Friday)",
                    "Rebrand it completely",
                    "Other (explain below)",
                ],
                key="loc_ex2_bf_choice",
            )
            bf_notes = st.text_area(
                "Explain your choice regarding **Black Friday**:",
                key="loc_ex2_bf_notes",
                height=100,
            )

            season_notes = st.text_area(
                "How did you handle the **seasonal / winter** reference?",
                key="loc_ex2_season_notes",
                height=100,
            )

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / prompts"):
            st.markdown(
                "- Visibility of Western commercial culture vs local practices\n"
                "- Relevance of winter imagery in Gulf advertising\n"
                "- Domestication vs foreignisation in marketing contexts"
            )

        if submitted:
            reflection = (
                "Literal translation:\n" + literal.strip() + "\n\n"
                "Decision on Black Friday:\n" + bf_choice + "\n" + bf_notes.strip() + "\n\n"
                "Seasonal adaptation notes:\n" + season_notes.strip()
            )
            save_loc_submission(ex_id, source_text, gulf, reflection)

    def exercise_3():
        ex_id = "LOC_3"
        items = [
            "The conference starts on 03/12/2026 at 9:00 AM.",
            "The package weighs 5 pounds and measures 12 inches.",
            "Prices are listed in USD.",
            "Submit your résumé before October 1.",
        ]
        source_text = "\n".join(items)

        st.header("3️⃣ Conventions: Dates, Units & Currency")
        st.write("Translate and localise the following items into Arabic:")

        with st.form(f"loc_form_{ex_id}"):
            answers = []
            for i, item in enumerate(items, start=1):
                st.markdown(f"**{i}. {item}**")
                ans = st.text_area(
                    f"Your localised Arabic version for item {i}:",
                    key=f"loc_ex3_item_{i}",
                    height=80,
                )
                answers.append(ans)

            st.markdown("### Reflection")
            amb = st.text_area(
                "Where could **ambiguity or misunderstanding** arise if conventions are not localised properly?",
                key="loc_ex3_ambiguity",
                height=120,
            )

            risk = st.text_area(
                "What are the **practical risks** (e.g., legal, financial, usability) of not localising these elements?",
                key="loc_ex3_risk",
                height=120,
            )

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / hints"):
            st.markdown(
                "- Date format ambiguity (03/12 vs 12/03)\n"
                "- Metric vs imperial units\n"
                "- Currency conversion and symbol localisation\n"
                "- CV vs résumé vs سيرة ذاتية"
            )

        if submitted:
            main_text = "\n".join(answers)
            reflection = (
                "Ambiguity notes:\n" + amb.strip() + "\n\n"
                "Risk notes:\n" + risk.strip()
            )
            save_loc_submission(ex_id, source_text, main_text, reflection)

    def exercise_4():
        ex_id = "LOC_4"
        source_text = "Welcome back, Sarah! We missed you. Ready to pick up where you left off?"

        st.header("4️⃣ Tone & Website/App UX")

        with st.form(f"loc_form_{ex_id}"):
            st.subheader("Source Text (English → Arabic)")
            st.markdown(f"> {source_text}")

            st.markdown("### Step 1 – Neutral MSA Translation")
            neutral = st.text_area(
                "Write a **neutral Modern Standard Arabic** translation:",
                key="loc_ex4_neutral",
                height=120,
            )

            st.markdown("### Step 2 – Contextual Localisation")

            col1, col2 = st.columns(2)
            with col1:
                gov = st.text_area(
                    "Localise for a **formal government portal**:",
                    key="loc_ex4_gov",
                    height=140,
                )
            with col2:
                fashion = st.text_area(
                    "Localise for a **fashion e-commerce website**:",
                    key="loc_ex4_fashion",
                    height=140,
                )

            st.markdown("### Step 3 – Tone & Trust")
            tone_trust = st.text_area(
                "How do **tone** and **register** affect user trust and engagement in each context?",
                key="loc_ex4_tone_trust",
                height=140,
            )

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / prompts"):
            st.markdown(
                "- Use of vocatives and name forms in Arabic\n"
                "- Degree of warmth vs distance in institutional voices\n"
                "- Second person pronoun choices (singular/plural, gender)"
            )

        if submitted:
            reflection = (
                "Neutral MSA version:\n" + neutral.strip() + "\n\n"
                "Government vs fashion tone notes:\n" + tone_trust.strip()
            )
            # Use the fashion e-commerce version as main text
            save_loc_submission(ex_id, source_text, fashion, reflection)

    def exercise_5():
        ex_id = "LOC_5"
        source_text = (
            "احصل على أفضل صفقات الجمعة السوداء الآن! الشحن مجاني لكل الطلبات أكثر من 50 دولار. "
            "اتصل بنا على 1-800-555-0199."
        )

        st.header("5️⃣ Post-editing: Error Detection & Localisation")

        with st.form(f"loc_form_{ex_id}"):
            st.subheader("Poorly Localised Arabic Banner")
            st.markdown(f"> {source_text}")

            st.markdown("### Step 1 – Identify Problems")
            issues = st.text_area(
                "List the **issues** you can spot:",
                key="loc_ex5_issues",
                height=140,
            )

            st.markdown("### Step 2 – Revised Version for UAE Market")
            revised = st.text_area(
                "Write an improved, **fully localised version for the UAE**:",
                key="loc_ex5_revised",
                height=160,
            )

            st.markdown("### Step 3 – Error Classification")
            col1, col2 = st.columns(2)
            with col1:
                ling = st.text_area("**Linguistic errors**", key="loc_ex5_ling", height=100)
                cult = st.text_area("**Cultural errors**", key="loc_ex5_cult", height=100)
            with col2:
                func = st.text_area("**Functional/technical errors**", key="loc_ex5_func", height=100)
                tech = st.text_area("**Formatting/numbering/other**", key="loc_ex5_tech", height=100)

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / prompts"):
            st.markdown(
                "- Appropriateness of “الجمعة السوداء” vs alternatives\n"
                "- Currency choice (dollars vs dirhams)\n"
                "- Phone number format and localisation\n"
                "- Register and marketing style in Arabic"
            )

        if submitted:
            reflection = (
                "Issues spotted:\n" + issues.strip() + "\n\n"
                "Linguistic errors:\n" + ling.strip() + "\n\n"
                "Cultural errors:\n" + cult.strip() + "\n\n"
                "Functional/technical errors:\n" + func.strip() + "\n\n"
                "Formatting/other:\n" + tech.strip()
            )
            save_loc_submission(ex_id, source_text, revised, reflection)

    def exercise_6():
        ex_id = "LOC_6"
        source_text = "Track your calories, crush your goals, and stay summer-ready all year long!"

        st.header("6️⃣ App Store Description Localisation")

        with st.form(f"loc_form_{ex_id}"):
            st.subheader("Source Text (English → Arabic)")
            st.markdown(f"> {source_text}")

            st.markdown("### Step 1 – Base Translation")
            base = st.text_area(
                "Write a **base translation into Arabic**:",
                key="loc_ex6_base",
                height=120,
            )

            st.markdown("### Step 2 – Variant A: Conservative Audience")
            conservative = st.text_area(
                "Localise for a more **conservative audience** (focus on health, well-being, moderation):",
                key="loc_ex6_conservative",
                height=140,
            )

            st.markdown("### Step 3 – Variant B: Youth-focused Fitness App")
            youth = st.text_area(
                "Localise for a **youth-oriented fitness app** (energetic, motivational tone):",
                key="loc_ex6_youth",
                height=140,
            )

            st.markdown("### Step 4 – Strategic Discussion")
            reg_notes = st.text_area(
                "How did you adapt **register**, **imagery**, and **implicit values** in each version?",
                key="loc_ex6_register",
                height=140,
            )

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / prompts"):
            st.markdown(
                "- Body image vs health framing\n"
                "- Use of motivational vs neutral language\n"
                "- Sensitivity to cultural norms around appearance"
            )

        if submitted:
            reflection = (
                "Base translation:\n" + base.strip() + "\n\n"
                "Register/imagery/values notes:\n" + reg_notes.strip()
            )
            # Use youth-focused version as main text
            save_loc_submission(ex_id, source_text, youth, reflection)

    def exercise_7():
        ex_id = "LOC_7"
        source_text = (
            "Two versions of the same promotional text: Version A (literal) "
            "and Version B (heavily localised)."
        )

        st.header("7️⃣ Strategy & Theory Reflection")

        with st.form(f"loc_form_{ex_id}"):
            st.write(
                "Imagine you have two Arabic versions of the same English promotional text:\n"
                "- **Version A:** literal translation\n"
                "- **Version B:** heavily localised adaptation\n\n"
                "You can either paste real examples below or answer hypothetically."
            )

            st.markdown("### Step 1 – (Optional) Paste Texts")
            ver_a = st.text_area("Paste **Version A (literal)** here:", key="loc_ex7_a", height=140)
            ver_b = st.text_area("Paste **Version B (localised)** here:", key="loc_ex7_b", height=140)

            st.markdown("### Step 2 – Skopos & Effect")
            skopos = st.text_area(
                "Which version better fulfils the **Skopos** (purpose) of the text, and why?",
                key="loc_ex7_skopos",
                height=140,
            )

            st.markdown("### Step 3 – Domestication vs Foreignisation")
            dom_for = st.text_area(
                "Where do you see **domestication** and **foreignisation** in the localised version?",
                key="loc_ex7_dom_for",
                height=140,
            )

            st.markdown("### Step 4 – Limits of Localisation")
            limits = st.text_area(
                "When does localisation risk becoming **over-adaptation** or even **rewriting**? Give examples.",
                key="loc_ex7_limits",
                height=140,
            )

            submitted = st.form_submit_button("Save & get feedback")

        with st.expander("👩‍🏫 Instructor notes / prompts"):
            st.markdown(
                "- Link discussion to Skopos theory\n"
                "- Discuss ethical and professional limits to adaptation\n"
                "- Consider client brief and audience expectations"
            )

        if submitted:
            main_text = ver_b if ver_b.strip() else ver_a
            reflection = (
                "Skopos analysis:\n" + skopos.strip() + "\n\n"
                "Domestication/foreignisation notes:\n" + dom_for.strip() + "\n\n"
                "Limits of localisation:\n" + limits.strip()
            )
            save_loc_submission(ex_id, source_text, main_text, reflection)

    # ---- Sticker/text/image-based tasks from instructor ----
    def exercise_stickers():
        ex_id_base = "LOC_STICKER"
        loc_stickers = load_json(LOC_STICKERS_FILE)

        if not loc_stickers:
            st.info("No sticker/text/image tasks have been created yet. Please ask your instructor.")
            return

        options = sorted(loc_stickers.keys())
        chosen_id = st.selectbox(
            "Choose task from your instructor",
            options,
            format_func=lambda sid: loc_stickers[sid].get("title", sid),
            key="loc_sticker_student_select"
        )

        data = loc_stickers.get(chosen_id, {})
        st.header("🎨 Sticker / text / image localisation task")
        st.subheader(data.get("title", chosen_id))

        content_text = data.get("content_text", "")
        if content_text:
            st.markdown("**Text to localise:**")
            st.write(content_text)

        # show image if any
        if data.get("image_type") == "uploaded" and data.get("image_path"):
            if Path(data["image_path"]).exists():
                st.image(data["image_path"])
            else:
                st.warning("Sticker/image file missing on server.")
        elif data.get("image_type") == "url" and data.get("image_url"):
            st.image(data["image_url"])

        if data.get("instructions"):
            st.markdown("**Instructor instructions:**")
            st.write(data["instructions"])

        with st.form(f"loc_form_{ex_id_base}_{chosen_id}"):
            answer = st.text_area(
                "Write your localised version in Arabic:",
                height=180,
                key=f"loc_stk_answer_{chosen_id}"
            )
            reflection = st.text_area(
                "Briefly explain your key localisation decisions:",
                height=120,
                key=f"loc_stk_refl_{chosen_id}"
            )
            submitted = st.form_submit_button("Submit task & get feedback")

        if submitted:
            # For metrics, treat source as combination of text + instructions
            source_text = "\n".join(
                [part for part in [content_text, data.get("instructions", "")] if part.strip()]
            )
            if not source_text.strip():
                source_text = data.get("title", chosen_id)
            ex_id = f"{ex_id_base}_{chosen_id}"
            save_loc_submission(ex_id, source_text, answer, reflection)

    # ---- Router ----
    ex_id_map = {
        "1️⃣ Translation vs Localisation": "LOC_1",
        "2️⃣ Cultural Adaptation in Advertising": "LOC_2",
        "3️⃣ Conventions: Dates, Units, Currency": "LOC_3",
        "4️⃣ Tone & Website/App UX": "LOC_4",
        "5️⃣ Post-editing: Error Detection": "LOC_5",
        "6️⃣ App Store Description": "LOC_6",
        "7️⃣ Strategy & Theory Reflection": "LOC_7",
        "🎨 Sticker / text / image task (from instructor)": "LOC_STICKER",
    }
    current_ex_id = ex_id_map.get(exercise, "LOC_STICKER")
    start_key = f"loc_start_{student_name}_{current_ex_id}"
    if start_key not in st.session_state:
        st.session_state[start_key] = time.time()

    if exercise == "1️⃣ Translation vs Localisation":
        exercise_1()
    elif exercise == "2️⃣ Cultural Adaptation in Advertising":
        exercise_2()
    elif exercise == "3️⃣ Conventions: Dates, Units, Currency":
        exercise_3()
    elif exercise == "4️⃣ Tone & Website/App UX":
        exercise_4()
    elif exercise == "5️⃣ Post-editing: Error Detection":
        exercise_5()
    elif exercise == "6️⃣ App Store Description":
        exercise_6()
    elif exercise == "7️⃣ Strategy & Theory Reflection":
        exercise_7()
    elif exercise == "🎨 Sticker / text / image task (from instructor)":
        exercise_stickers()

# ---------------- Main ----------------
def main():
    st.set_page_config(page_title="Translation Lab (EduApp)", layout="wide")

    st.write("HF token loaded:", bool(os.getenv("HF_API_TOKEN")))

    st.sidebar.title("Navigation")

    st.markdown(
        "<div style='padding:8px;border:1px solid #ddd;border-radius:8px;background:#f7f9ff'>"
        "<b>EduApp – Build:</b> 2025-11-10 v5 (translation + localisation lab, grant-aligned)</div>",
        unsafe_allow_html=True
    )

    section = st.sidebar.radio(
        "Module",
        ["Core Translation Lab", "Localisation Lab"],
        index=0
    )

    if section == "Localisation Lab":
        localisation_lab()
        return

    # Core translation lab
    role = st.sidebar.radio("Login as", ["Instructor", "Student"], index=1)
    if role == "Instructor":
        instructor_dashboard()
    else:
        student_dashboard()

if __name__ == "__main__":
    main()
